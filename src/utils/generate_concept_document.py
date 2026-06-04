"""
generate_concept_document.py — Create professional DOCX explaining SENTINEL-GNSS concepts.

Generates a clean, professional academic document with:
- Cividis color palette (no dark backgrounds)
- Bold text for emphasis
- Clear section structure
- Scenario explanations
- Suitable for stakeholder review
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Cividis color palette (RGB values)
COLORS = {
    "c0": (0, 32, 96),         # Dark blue
    "c1": (52, 80, 127),       # Blue
    "c2": (91, 119, 140),      # Blue-grey
    "c3": (127, 155, 131),     # Teal-grey
    "c4": (188, 178, 69),      # Yellow-green
    "c5": (254, 231, 92),      # Yellow
    "black": (0, 0, 0),
    "grey": (122, 122, 122),
    "white": (255, 255, 255),
}

# Font sizes
SIZES = {
    "title": 18,
    "heading": 14,
    "subheading": 12,
    "body": 11,
    "small": 10,
}

def add_heading(doc, text, level=1, color=None):
    """Add a heading with optional color."""
    if color is None:
        color = COLORS["c0"]  # Default: dark blue

    h = doc.add_paragraph()
    h.style = f"Heading {level}"
    r = h.add_run(text)
    r.font.size = Pt(SIZES["heading"] if level == 1 else SIZES["subheading"])
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_body_text(doc, text, bold=False, color=None):
    """Add body text with optional bold and color."""
    if color is None:
        color = COLORS["black"]

    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    for run in p.runs:
        run.font.size = Pt(SIZES["body"])
        run.font.bold = bold
        run.font.color.rgb = RGBColor(*color)

    return p

def add_bold_text(doc, text, color=None):
    """Add bold text."""
    if color is None:
        color = COLORS["c1"]
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(SIZES["body"])
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(6)
    return p

def generate_document():
    """Generate the concept document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(SIZES["body"])

    # ---- SECTION 1: WHAT IS GNSS? ----
    add_heading(doc, "What is GNSS and Why It Matters", level=1, color=COLORS["c0"])

    add_body_text(doc, "Global Navigation Satellite System (GNSS) is how vehicles know their precise position. In most places, it works perfectly. But in cities, GNSS fails suddenly:")

    p = doc.add_paragraph()
    for item in [
        "Urban canyons: tall buildings block signals",
        "Tunnels: no line-of-sight to satellites",
        "Foliage: tree canopies scatter signals",
        "Multipath: signals bounce off buildings, creating errors"
    ]:
        p.add_run(f"• {item}\n").font.size = Pt(SIZES["body"])
    p.paragraph_format.space_after = Pt(6)

    add_body_text(doc, "Current systems only notice failure AFTER it happens. SENTINEL-GNSS predicts it 5–30 seconds ahead.")

    # ---- SECTION 2: THREE SIGNAL CLASSES ----
    add_heading(doc, "Three Signal Quality Classes", level=1, color=COLORS["c0"])

    # Class 1: CLEAN
    add_bold_text(doc, "[CLEAN] Healthy Signal", color=COLORS["c0"])
    add_body_text(doc, "Full constellation of satellites visible. High C/N0 (signal strength ~40+ dB-Hz). Excellent position accuracy. No degradation risk.", color=COLORS["black"])
    add_body_text(doc, "Examples: open sky, plaza, highway", color=COLORS["grey"])
    doc.add_paragraph()

    # Class 2: WARNING
    add_bold_text(doc, "[WARNING] Partial Degradation", color=COLORS["c4"])
    add_body_text(doc, "Some satellites lost due to local blockage. C/N0 degrading (30-40 dB-Hz). Position accuracy declining. Risk of further degradation.", color=COLORS["black"])
    add_body_text(doc, "Examples: urban street with tall buildings, partial tree canopy", color=COLORS["grey"])
    doc.add_paragraph()

    # Class 3: DEGRADED
    add_bold_text(doc, "[DEGRADED] Severe Loss or No Fix", color=COLORS["c5"])
    add_body_text(doc, "Few/no satellites, very low C/N₀ (<30 dB-Hz), or complete loss of fix. Position is unreliable or unavailable. Vehicle must switch to backup localization (IMU, map-matching).", color=COLORS["black"])
    add_body_text(doc, "Examples: dense urban canyon, tunnel, heavy tree cover", color=COLORS["grey"])
    doc.add_paragraph()

    # ---- SECTION 3: FIVE SCENARIOS ----
    add_heading(doc, "Collection Scenarios (A–E)", level=1, color=COLORS["c0"])

    scenarios = [
        ("Scenario A: Instant Blockage",
         "Sharp, sudden transition from CLEAN to DEGRADED. Example: driving under a bridge. Tests the model's ability to catch rapid changes."),
        ("Scenario B: Urban Canyon",
         "Gradual signal loss while moving between tall buildings. Multipath errors increase. Tests degradation trend detection."),
        ("Scenario C: Partial Blockage",
         "Stable but reduced signal under trees or partial cover. Some satellites visible, but not enough for optimal position. Tests WARNING detection."),
        ("Scenario D: Open Sky",
         "Clean baseline with strongest signal and best geometry. Tests that model correctly identifies CLEAN periods."),
        ("Scenario E: Approaching Blockage",
         "Smooth signal degradation while moving toward a blocking structure. Tests the model's ability to predict an approaching failure."),
    ]

    for title, desc in scenarios:
        add_bold_text(doc, title, color=COLORS["c1"])
        add_body_text(doc, desc, color=COLORS["black"])
        doc.add_paragraph()

    # ---- SECTION 4: THE PREDICTION MODEL ----
    add_heading(doc, "How SENTINEL-GNSS Predicts Degradation", level=1, color=COLORS["c0"])

    add_body_text(doc, "The model reads 30 seconds of signal history and predicts:")

    p = doc.add_paragraph()
    for item in ["P(CLEAN) at +5s / +15s / +30s", "P(WARNING) at +5s / +15s / +30s", "P(DEGRADED) at +5s / +15s / +30s"]:
        r = p.add_run(f"• {item}\n")
        r.font.size = Pt(SIZES["body"])
        r.font.bold = True
        r.font.color.rgb = RGBColor(*COLORS["c1"])
    p.paragraph_format.space_after = Pt(6)

    add_body_text(doc, "Architecture:", bold=True, color=COLORS["c1"])
    add_body_text(doc, "Transformer encoder (sees long-range patterns in signal history) → BiLSTM (captures degradation trends) → Three output heads (one per horizon).")

    add_body_text(doc, "37 engineered features from raw GNSS observations:", bold=True, color=COLORS["c1"])
    features_text = "Signal strength (C/N₀), satellite geometry (DOP), constellation count, receiver status, temporal patterns (how C/N₀ is changing), atmospheric effects."
    add_body_text(doc, features_text, color=COLORS["black"])

    # ---- SECTION 5: VALIDATION ----
    add_heading(doc, "How We Validate (Cross-City Proof)", level=1, color=COLORS["c0"])

    add_body_text(doc, "We train on Beihang (Beijing) and Hong Kong data. Then we test on Tokyo—a city the model has never seen.")

    add_bold_text(doc, "Why is this important?", color=COLORS["c1"])
    add_body_text(doc, "A model trained in Beijing might fail in Tokyo if it only learned Beijing-specific patterns. Predicting well on an unseen city proves true generalization.")

    add_body_text(doc, "Results: 89% accuracy on Tokyo, including 90% on the safety-critical DEGRADED class.", bold=True, color=COLORS["c1"])

    # ---- SECTION 6: EKF ----
    add_heading(doc, "Using Prediction to Improve Navigation: The Adaptive EKF", level=1, color=COLORS["c0"])

    add_body_text(doc, "Prediction is useful only if we use it. The Extended Kalman Filter fuses GNSS measurements with a motion model to estimate true position. Our innovation:")

    add_bold_text(doc, "Adaptive Measurement Noise", color=COLORS["c1"])
    add_body_text(doc, "When P(DEGRADED) is high → inflate measurement noise (distrust GNSS) → lean on motion model (dead-reckoning)")
    add_body_text(doc, "When P(DEGRADED) is low → normal measurement noise → use GNSS to correct drift", color=COLORS["black"])

    add_body_text(doc, "Result: 34% improvement in position accuracy during blockage (synthetic test).", bold=True, color=COLORS["c4"])

    # ---- SECTION 7: REAL-WORLD VALIDATION ----
    add_heading(doc, "Phase 2a: Real-World Validation on UrbanNav Tokyo", level=1, color=COLORS["c0"])

    add_body_text(doc, "The synthetic blockage proves the concept. For real-world proof, we use UrbanNav Tokyo dataset because:")

    p = doc.add_paragraph()
    items = [
        "[OK] Has cm-level ground truth (SPAN-INS post-processed trajectory)",
        "[OK] Real GNSS observations with actual blockage events",
        "[OK] IMU data (accelerometer + gyro) for sensor fusion",
        "[OK] Public dataset, reproducible, peer-reviewed"
    ]
    for item in items:
        r = p.add_run(f"{item}\n")
        r.font.size = Pt(SIZES["body"])
        r.font.bold = False
        r.font.color.rgb = RGBColor(*COLORS["c0"])
    p.paragraph_format.space_after = Pt(6)

    add_body_text(doc, "Expected outcome: 15–30% RMSE improvement during actual urban blockage.", bold=True, color=COLORS["c1"])

    # ---- FOOTER ----
    doc.add_paragraph()
    footer = doc.add_paragraph("SENTINEL-GNSS Project | Beihang University, Beijing | 2026")
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(SIZES["small"])
        run.font.color.rgb = RGBColor(*COLORS["grey"])

    # Save
    output_path = Path(__file__).resolve().parents[2] / "SENTINEL_GNSS_Concepts.docx"
    doc.save(output_path)
    print(f"[OK] Document saved: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_document()
